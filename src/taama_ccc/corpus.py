from __future__ import annotations

import re
from hashlib import sha256
from pathlib import Path

import docx
from docx.document import Document as DocxDocumentType
from docx.oxml.table import CT_Tbl
from docx.oxml.text.paragraph import CT_P
from docx.table import Table, _Cell
from docx.text.paragraph import Paragraph

from taama_ccc.models import Document, DocumentChunk, DocumentSource

STALE_MARKERS = [
    "removed in",
    "no longer available",
    "superseded",
    "outdated",
    "refer to v20",
    "check updated",
    "unable to verify",
    "unable to provide",
    "not done",
    "not available",
    "alt form",
]

_CHECKBOX_MARK_RE = re.compile(r"(?<![☐xX✗])[xX✗]\s*([A-Za-z][A-Za-z /\-]*)")
_HEADING_STYLES = {"Heading 1", "Heading 2", "Heading 3"}


def _iter_block_items(parent: DocxDocumentType):
    for child in parent.element.body.iterchildren():
        if isinstance(child, CT_P):
            yield "p", Paragraph(child, parent)
        elif isinstance(child, CT_Tbl):
            yield "tbl", Table(child, parent)


def _parse_checkbox(text: str) -> str | None:
    marked = _CHECKBOX_MARK_RE.findall(text)

    return marked[0].strip().rstrip(".") if marked else None


def _detect_staleness(*texts: str) -> str | None:
    joined = " ".join(texts).lower()

    for marker in STALE_MARKERS:
        if marker in joined:
            return marker

    return None


def _extract_cell_links(cell: _Cell) -> list[str]:
    links: list[str] = []

    for paragraph in cell.paragraphs:
        for hyperlink in paragraph._p.xpath(".//w:hyperlink"):
            rel_id = hyperlink.get(
                "{http://schemas.openxmlformats.org/officeDocument/2006/relationships}id"
            )

            if rel_id is None:
                continue

            rel = cell.part.rels.get(rel_id)

            if rel is not None:
                links.append(rel.target_ref)

    return list(dict.fromkeys(links))


def _chunk_id(*parts: str) -> str:
    return sha256(":".join(parts).encode("utf-8")).hexdigest()


def parse_docx_table_rows(path: Path) -> list[DocumentChunk]:
    doc = docx.Document(str(path))
    document_id = path.stem

    current_section: str | None = None
    current_category: str | None = None
    chunks: list[DocumentChunk] = []
    row_index = 0

    for kind, obj in _iter_block_items(doc):
        if kind == "p":
            style_name = obj.style.name if obj.style else None

            if style_name in _HEADING_STYLES and obj.text.strip():
                current_section = obj.text.strip()

            continue

        table: Table = obj
        nrows, ncols = len(table.rows), len(table.columns)

        if nrows == 1 and ncols == 1:
            title = table.rows[0].cells[0].text.strip()

            if title and title.lower() != "nil":
                current_category = title

            continue

        header_cells = [c.text.strip() for c in table.rows[0].cells]

        if not header_cells or all(len(h) < 2 for h in header_cells):
            continue

        for row in table.rows[1:]:
            cells = row.cells

            if len(cells) != len(header_cells):
                continue

            texts = [c.text.strip() for c in cells]
            non_empty = [t for t in texts if t]

            if len(non_empty) == 1 and texts[0] == non_empty[0]:
                current_category = f"{current_category or ''} / {non_empty[0]}".strip(
                    " /"
                )
                continue

            field_map = dict(zip(header_cells, texts, strict=True))
            requirement_area_key = (
                next(
                    (
                        key
                        for key in (
                            "Requirement area",
                            "Criterion",
                            "Trigger",
                            "List / database",
                            "Condition type",
                            "Link type",
                            "Question",
                            "Question for the consultant",
                            "Statement",
                        )
                        if field_map.get(key)
                    ),
                    header_cells[0],
                ),
            )
            requirement_area = field_map.get(requirement_area_key, header_cells[0])
            comment = (
                field_map.get("Comment")
                or field_map.get("Answer / correction")
                or field_map.get("Answer / finding")
                or field_map.get("Consultant answer")
                or ""
            )
            approve_field = next(
                (
                    v
                    for k, v in field_map.items()
                    if any(term in k.lower() for term in ("approve", "correct", "ok?"))
                ),
                "",
            )

            links: list[str] = []

            for cell in cells:
                links.extend(_extract_cell_links(cell))

            links = list(dict.fromkeys(links))

            stale_marker = _detect_staleness(comment, *texts)
            approve_status = _parse_checkbox(approve_field)

            row_index += 1
            chunk_id = _chunk_id(
                document_id,
                str(current_section),
                str(current_category),
                str(row_index),
                requirement_area,
            )

            text_lines = [
                f"[Section: {current_section or '—'} | Category: {current_category or '—'}]",
                f"Requirement/Row: {requirement_area}",
            ]

            for key, value in field_map.items():
                if key in (header_cells[0], requirement_area_key) or not value:
                    continue

                text_lines.append(f"{key}: {value}")

            if links:
                text_lines.append(f"Sources: {', '.join(links)}")

            if stale_marker:
                text_lines.append(
                    f"[FLAG: possible stale/superseded source — matched '{stale_marker}']"
                )

            # Consultant's illustrative sample product set
            is_example = (current_section or "").lower().startswith("part 4")

            chunks.append(
                DocumentChunk(
                    id=chunk_id,
                    document_id=document_id,
                    text="\n".join(text_lines),
                    section=current_section,
                    source_url=links[0] if links else None,
                    metadata={
                        "chunk_type": "table_row",
                        "document_section": "illustrative_example"
                        if is_example
                        else "binding_rule",
                        "category": current_category or "",
                        "requirement_area": requirement_area,
                        "approve_status": approve_status or "",
                        "comment": comment,
                        "source_links": ", ".join(links),
                        "possible_stale_source": "true" if stale_marker else "false",
                        "stale_marker": stale_marker or "",
                    },
                )
            )

    return chunks


def load_docx_prose(path: Path) -> Document:
    doc = docx.Document(str(path))
    paragraphs: list[str] = []

    for kind, obj in _iter_block_items(doc):
        if kind != "p":
            continue

        style_name = obj.style.name if obj.style else None

        if style_name in _HEADING_STYLES:
            continue

        text = obj.text.strip()

        if text:
            paragraphs.append(text)

    return Document(
        id=path.stem,
        source=DocumentSource.DOCX,
        title=path.stem,
        content="\n\n".join(paragraphs),
    )


def chunk_document(document: Document, *, max_chars: int = 4000) -> list[DocumentChunk]:
    paragraphs = [p for p in document.content.split("\n\n") if p.strip()]

    chunks: list[DocumentChunk] = []
    buffer: list[str] = []
    buffer_len = 0
    index = 0

    def flush() -> None:
        nonlocal buffer, buffer_len, index

        if not buffer:
            return

        text = "\n\n".join(buffer)

        chunks.append(
            DocumentChunk(
                id=_chunk_id(document.id, "prose", str(index)),
                document_id=document.id,
                text=text,
                metadata={"chunk_type": "prose"},
            )
        )

        index += 1
        buffer = []
        buffer_len = 0

    for paragraph in paragraphs:
        if buffer and buffer_len + len(paragraph) > max_chars:
            flush()

        buffer.append(paragraph)
        buffer_len += len(paragraph)

    flush()

    return chunks


def parse_corpus(path: Path) -> list[DocumentChunk]:
    chunks = list(parse_docx_table_rows(path))
    chunks.extend(chunk_document(load_docx_prose(path)))

    return chunks
